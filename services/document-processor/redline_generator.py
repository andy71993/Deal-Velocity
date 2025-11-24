from docx import Document
from docx.shared import RGBColor
import io
import difflib
import re

def normalize_text(text):
    """Normalize whitespace for better matching"""
    return re.sub(r'\s+', ' ', text).strip()

def find_fuzzy_location(haystack, needle, threshold=0.8):
    """
    Finds the best match of needle in haystack using SequenceMatcher.
    Returns (start_index, end_index) in haystack, or (None, None).
    """
    if not needle or not haystack:
        return None, None
        
    # 1. Try exact match first (fastest)
    try:
        idx = haystack.index(needle)
        return idx, idx + len(needle)
    except ValueError:
        pass
        
    # 2. Try normalized match (if whitespace differs)
    # This is hard to map back indices, so we skip to fuzzy
    
    # 3. Fuzzy match using SequenceMatcher
    # This finds the longest contiguous matching block
    matcher = difflib.SequenceMatcher(None, haystack, needle)
    match = matcher.find_longest_match(0, len(haystack), 0, len(needle))
    
    # Check if the match is good enough
    if match.size >= len(needle) * threshold:
        return match.a, match.a + match.size
        
    # If the longest match isn't enough, it might be fragmented.
    # For MVP, we'll assume the clause is mostly contiguous.
    # If we really can't find it, we return None.
    return None, None

def generate_diff_ops(original_text, final_text):
    """
    Generates a list of operations (insert, delete, equal) to transform original to final.
    """
    matcher = difflib.SequenceMatcher(None, original_text, final_text)
    return matcher.get_opcodes()

def create_redlined_document(original_text: str, changes: list):
    """
    Generates a .docx file with track changes.
    
    Args:
        original_text (str): The full original contract text.
        changes (list): List of dicts with 'original' and 'new' text.
    """
    # 1. Apply changes to create the "Final" text
    # We do this by finding locations and replacing
    
    replacements = [] # (start, end, new_text)
    
    for change in changes:
        orig = change.get('original', '')
        new = change.get('new', '')
        
        if not orig or not new:
            continue
            
        start, end = find_fuzzy_location(original_text, orig)
        if start is not None:
            replacements.append((start, end, new))
        else:
            print(f"Warning: Could not find clause in text: {orig[:30]}...")
            
    # Sort replacements by start index descending to apply them without shifting indices
    replacements.sort(key=lambda x: x[0], reverse=True)
    
    final_text = original_text
    for start, end, new in replacements:
        final_text = final_text[:start] + new + final_text[end:]
        
    # 2. Generate diff between Original and Final
    # We use word-level splitting for better diffs than character-level
    # But SequenceMatcher works on any sequence. Let's try character level first 
    # as it preserves formatting best, but word level is more readable for redlines.
    
    # Let's try a hybrid: Split by lines or words?
    # Character level diffs can be messy (e.g. "th[e]->[a]t").
    # Word level is standard for redlines.
    
    def tokenize(text):
        return re.split(r'(\s+)', text)
        
    orig_tokens = tokenize(original_text)
    final_tokens = tokenize(final_text)
    
    matcher = difflib.SequenceMatcher(None, orig_tokens, final_tokens)
    
    doc = Document()
    para = doc.add_paragraph()
    
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            # Text is same
            text = "".join(orig_tokens[i1:i2])
            _add_text_run(para, text)
        elif tag == 'replace':
            # Delete original
            del_text = "".join(orig_tokens[i1:i2])
            _add_text_run(para, del_text, is_deleted=True)
            # Insert new
            ins_text = "".join(final_tokens[j1:j2])
            _add_text_run(para, ins_text, is_inserted=True)
        elif tag == 'delete':
            del_text = "".join(orig_tokens[i1:i2])
            _add_text_run(para, del_text, is_deleted=True)
        elif tag == 'insert':
            ins_text = "".join(final_tokens[j1:j2])
            _add_text_run(para, ins_text, is_inserted=True)
            
    # Save to memory
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def _add_text_run(para, text, is_deleted=False, is_inserted=False):
    """Helper to add text with formatting, handling newlines"""
    if not text:
        return
        
    # Handle newlines - docx runs don't handle \n well, usually need new paragraph
    # But we are inside one paragraph. 
    # If text contains newlines, we might need to split?
    # For simplicity in this MVP, we'll replace \n with a break or just let it be.
    # Actually, docx supports .add_break() but that's for runs.
    
    # Simple approach: just add the run. 
    # If it looks bad, we can improve newline handling later.
    
    run = para.add_run(text)
    if is_deleted:
        run.font.strike = True
        run.font.color.rgb = RGBColor(255, 0, 0) # Red
    if is_inserted:
        run.font.underline = True
        run.font.color.rgb = RGBColor(0, 0, 255) # Blue
